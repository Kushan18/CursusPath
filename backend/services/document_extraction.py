"""
Extracts plain text from an uploaded offer letter or resume.

Strategy:
- Native PDF with a real text layer -> PyMuPDF pulls the text directly (fast, exact).
- PDF that is actually a scanned image (no text layer) -> rasterize each page
  and run it through Tesseract OCR.
- DOCX -> python-docx pulls paragraph and table text directly.
- JPG/PNG upload -> straight to Tesseract OCR.
"""

import io
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document

MIN_TEXT_LENGTH_BEFORE_OCR_FALLBACK = 40


class ExtractionError(Exception):
    pass


def _extract_docx_text(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts).strip()


def _extract_pdf_text_layer(file_bytes: bytes) -> str:
    text_parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts).strip()


def _ocr_pdf_pages(file_bytes: bytes) -> str:
    try:
        text_parts = []
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                text_parts.append(pytesseract.image_to_string(img))
        return "\n".join(text_parts).strip()
    except (FileNotFoundError, pytesseract.TesseractNotFoundError):
        raise ExtractionError("Tesseract OCR is not installed on this server. Please upload a standard text-based PDF or DOCX instead of a scanned image.")
    except Exception as e:
        raise ExtractionError(f"OCR failed: {e}")

def _ocr_image(file_bytes: bytes) -> str:
    try:
        img = Image.open(io.BytesIO(file_bytes))
        return pytesseract.image_to_string(img).strip()
    except (FileNotFoundError, pytesseract.TesseractNotFoundError):
        raise ExtractionError("Tesseract OCR is not installed on this server. Please upload a standard text-based PDF or DOCX instead of an image.")
    except Exception as e:
        raise ExtractionError(f"OCR failed: {e}")


def extract_text(file_bytes: bytes, content_type: str, filename: str) -> str:
    is_pdf = content_type == "application/pdf" or filename.lower().endswith(".pdf")
    is_docx = (
        content_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or filename.lower().endswith(".docx")
    )
    is_image = content_type.startswith("image/") or filename.lower().endswith(
        (".png", ".jpg", ".jpeg", ".webp")
    )

    if is_pdf:
        text = _extract_pdf_text_layer(file_bytes)
        if len(text) < MIN_TEXT_LENGTH_BEFORE_OCR_FALLBACK:
            text = _ocr_pdf_pages(file_bytes)
    elif is_docx:
        text = _extract_docx_text(file_bytes)
    elif is_image:
        text = _ocr_image(file_bytes)
    else:
        raise ExtractionError(
            f"Unsupported file type '{content_type}'. Upload a PDF, DOCX, JPG, or PNG."
        )

    if len(text.strip()) < MIN_TEXT_LENGTH_BEFORE_OCR_FALLBACK:
        raise ExtractionError(
            "Could not extract readable text from this file. "
            "Try a clearer scan or a different file."
        )

    return text
