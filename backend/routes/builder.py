import logging
import uuid
import io
import fitz
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
from fastapi import APIRouter, Depends, UploadFile, File, Form, HTTPException
from fastapi.responses import StreamingResponse
from middleware.auth import get_user_and_scoped_client
from langfuse.decorators import observe, langfuse_context
from docx import Document
from services.resume_gemini_service import _call_gemini, GeminiAnalysisError
from services.resume_groq_service import _call_groq, GroqAnalysisError

logger = logging.getLogger("cursuspath.builder")
router = APIRouter(prefix="/api/v1/builder", tags=["builder"])


class ResumeState(BaseModel):
    resume_name: str
    target_role: Optional[str] = ""
    job_description: Optional[str] = ""
    resume_data: Dict[str, Any]
    skipped_fields: List[Dict[str, str]]
    template_id: str

class UpdateResumeState(BaseModel):
    resume_name: Optional[str] = None
    target_role: Optional[str] = None
    job_description: Optional[str] = None
    resume_data: Optional[Dict[str, Any]] = None
    skipped_fields: Optional[List[Dict[str, str]]] = None
    template_id: Optional[str] = None
    parseability_score: Optional[int] = None
    job_match_score: Optional[int] = None
    score_deductions: Optional[List[str]] = None

@router.get("/")
async def get_built_resumes(user_and_client=Depends(get_user_and_scoped_client)):
    current_user, supabase = user_and_client
    try:
        response = (
            supabase.table("user_resumes")
            .select("*")
            .eq("user_id", current_user.id)
            .order("updated_at", desc=True)
            .execute()
        )
        return response.data
    except Exception as e:
        logger.error("Failed to fetch user_resumes: %s", e)
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/{resume_id}")
async def get_built_resume(resume_id: str, user_and_client=Depends(get_user_and_scoped_client)):
    current_user, supabase = user_and_client
    try:
        response = (
            supabase.table("user_resumes")
            .select("*")
            .eq("id", resume_id)
            .eq("user_id", current_user.id)
            .execute()
        )
        if not response.data:
            raise HTTPException(status_code=404, detail="Resume not found")
        return response.data[0]
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Failed to fetch user_resume: %s", e)
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/")
async def save_built_resume(state: ResumeState, user_and_client=Depends(get_user_and_scoped_client)):
    current_user, supabase = user_and_client
    row = {
        "user_id": current_user.id,
        "resume_name": state.resume_name,
        "target_role": state.target_role,
        "job_description": state.job_description,
        "resume_data": state.resume_data,
        "skipped_fields": state.skipped_fields,
        "template_id": state.template_id,
        "parseability_score": 0,
        "job_match_score": 0,
        "score_deductions": []
    }
    try:
        response = supabase.table("user_resumes").insert(row).execute()
        return response.data[0] if response.data else row
    except Exception as e:
        logger.error("Failed to save user_resume: %s", e)
        raise HTTPException(status_code=500, detail=str(e))

@router.put("/{resume_id}")
async def update_built_resume(resume_id: str, state: UpdateResumeState, user_and_client=Depends(get_user_and_scoped_client)):
    current_user, supabase = user_and_client
    update_data = {"updated_at": "now()"}
    if state.resume_name is not None: update_data["resume_name"] = state.resume_name
    if state.target_role is not None: update_data["target_role"] = state.target_role
    if state.job_description is not None: update_data["job_description"] = state.job_description
    if state.resume_data is not None: update_data["resume_data"] = state.resume_data
    if state.skipped_fields is not None: update_data["skipped_fields"] = state.skipped_fields
    if state.template_id is not None: update_data["template_id"] = state.template_id
    if state.parseability_score is not None: update_data["parseability_score"] = state.parseability_score
    if state.job_match_score is not None: update_data["job_match_score"] = state.job_match_score
    if state.score_deductions is not None: update_data["score_deductions"] = state.score_deductions
    try:
        response = (
            supabase.table("user_resumes")
            .update(update_data)
            .eq("id", resume_id)
            .eq("user_id", current_user.id)
            .execute()
        )
        return response.data[0] if response.data else {"id": resume_id}
    except Exception as e:
        logger.error("Failed to update user_resume: %s", e)
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/{resume_id}/duplicate")
async def duplicate_built_resume(resume_id: str, user_and_client=Depends(get_user_and_scoped_client)):
    current_user, supabase = user_and_client
    try:
        existing = (
            supabase.table("user_resumes")
            .select("*")
            .eq("id", resume_id)
            .eq("user_id", current_user.id)
            .execute()
        )
        if not existing.data:
            raise HTTPException(status_code=404, detail="Resume not found")
        
        row = existing.data[0]
        del row["id"]
        del row["created_at"]
        del row["updated_at"]
        row["resume_name"] = row["resume_name"] + " (Copy)"
        
        inserted = supabase.table("user_resumes").insert(row).execute()
        return inserted.data[0] if inserted.data else row
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Failed to duplicate user_resume: %s", e)
        raise HTTPException(status_code=500, detail=str(e))

@router.delete("/{resume_id}")
async def delete_built_resume(resume_id: str, user_and_client=Depends(get_user_and_scoped_client)):
    current_user, supabase = user_and_client
    try:
        supabase.table("user_resumes").delete().eq("id", resume_id).eq("user_id", current_user.id).execute()
        return {"deleted": True, "id": resume_id}
    except Exception as e:
        logger.error("Failed to delete user_resume: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


class GenerateContentRequest(BaseModel):
    action: str  # "create_from_profile", "refine_resume", "rewrite_bullet", "generate_summary"
    content: str
    target_role: Optional[str] = ""
    job_description: Optional[str] = ""
    user_instruction: Optional[str] = ""

@observe(as_type="generation")
def _run_ai_generation(sys_prompt: str, user_prompt: str) -> dict:
    try:
        return _call_gemini(sys_prompt, user_prompt)
    except GeminiAnalysisError as e:
        logger.warning(f"Gemini failed for builder: {e}, falling back to Groq")
        try:
            return _call_groq(sys_prompt, user_prompt)
        except GroqAnalysisError as e2:
            raise AllProvidersFailedError("Both AI providers failed.") from e2

class AllProvidersFailedError(Exception):
    pass

@router.post("/generate")
async def generate_content(req: GenerateContentRequest, user_and_client=Depends(get_user_and_scoped_client)):
    current_user, _ = user_and_client
    langfuse_context.update_current_observation(user_id=current_user.id)
    
    sys_prompt = "You are an expert resume writer. "
    user_prompt = ""
    
    schema_str = (
        "{\n"
        '  "contact": { "fullName": "", "email": "", "phone": "", "location": "", "linkedin": "", "github": "" },\n'
        '  "summary": "",\n'
        '  "experience": [ { "title": "", "company": "", "dates": "", "location": "", "bullets": ["", ""] } ],\n'
        '  "education": [ { "degree": "", "school": "", "dates": "" } ],\n'
        '  "skills": ["", ""],\n'
        '  "projects": [ { "title": "", "company": "", "dates": "", "location": "", "bullets": ["", ""] } ],\n'
        '  "certifications": []\n'
        "}\n"
    )

    if req.action == "create_from_profile":
        sys_prompt += (
            "Construct a highly professional, 1-page ATS-optimized resume. "
            "Write exactly like a human expert—avoid obvious AI buzzwords like 'spearheaded', 'synergized', 'dynamic', or 'delved'. "
            "Focus strictly on concrete achievements, technical metrics, and clear impact. "
            f"Output MUST be ONLY a JSON object exactly matching this schema:\n{schema_str}"
        )
        user_prompt = f"PROFILE DATA:\n{req.content}"
    elif req.action == "refine_resume":
        sys_prompt += (
            "You are an AI resume building assistant. The user wants to change something about their resume or format. "
            "If the user asks to change the format/style/template, set 'show_template_options' to true. "
            "Apply any content changes intelligently while maintaining a professional, human tone. "
            "Output MUST be ONLY a JSON object exactly matching this schema:\n"
            "{\n"
            '  "resume_data": ' + schema_str.replace('\n', '') + ',\n'
            '  "chat_response": "A friendly message explaining what you changed",\n'
            '  "show_template_options": false\n'
            "}"
        )
        user_prompt = f"USER INSTRUCTION: {req.user_instruction}\n\nCURRENT RESUME:\n{req.content}"
    elif req.action == "generate_review":
        sys_prompt += (
            "You are an expert ATS Resume Reviewer. Review the provided resume data. "
            "Write a short, professional message to the user. "
            "Point out what looks correct/strong, and point out what is wrong or could be improved. "
            "Output MUST be ONLY a JSON object exactly matching this schema:\n"
            "{\n"
            '  "review": "Your review text here..."\n'
            "}"
        )
        user_prompt = f"CURRENT RESUME:\n{req.content}"
    elif req.action == "analyze_missing_fields":
        sys_prompt += (
            "You are an expert ATS Resume Reviewer. Review the provided resume data against the target role. "
            "Identify 1-3 critical missing elements (e.g., missing LinkedIn, missing a summary, missing specific certifications like AWS for cloud roles, or missing quantitative metrics). "
            "Output MUST be ONLY a JSON object exactly matching this schema:\n"
            "{\n"
            '  "missing_fields": [\n'
            '    { "field": "Name of field (e.g. LinkedIn)", "reason": "Professional explanation of why it is needed", "type": "input" }\n'
            '  ]\n'
            "}"
        )
        user_prompt = f"CURRENT RESUME:\n{req.content}"
    elif req.action == "rewrite_bullet":
        sys_prompt += "Rewrite the following resume bullet point to be achievement-oriented, using strong action verbs and quantifying results where implied. Do not invent new metrics."
        sys_prompt += "\nRespond with ONLY a JSON object: { \"suggestion\": \"...\" }"
        user_prompt = f"BULLET:\n{req.content}"
    elif req.action == "generate_summary":
        sys_prompt += "Write a professional summary for a resume based on the following background. Keep it to 2-3 sentences."
        sys_prompt += "\nRespond with ONLY a JSON object: { \"suggestion\": \"...\" }"
        user_prompt = f"BACKGROUND:\n{req.content}"
    else:
        raise HTTPException(status_code=400, detail="Invalid action")
        
    if req.target_role:
        user_prompt += f"\n\nTARGET ROLE: {req.target_role}"
    if req.job_description:
        user_prompt += f"\n\nJOB DESCRIPTION: {req.job_description}"
        
    try:
        result = _run_ai_generation(sys_prompt, user_prompt)
        langfuse_context.flush()
        return result
    except Exception as e:
        raise HTTPException(status_code=502, detail=str(e))


@router.post("/score")
async def score_resume(
    file: UploadFile = File(...),
    skipped_fields: str = Form("[]"),
    job_description: str = Form(""),
    user_and_client=Depends(get_user_and_scoped_client)
):
    import json
    try:
        skipped = json.loads(skipped_fields)
    except:
        skipped = []

    file_bytes = await file.read()
    
    # Parseability Score (Base 100)
    parseability_score = 100
    deductions = []
    improvements = []

    try:
        pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in pdf_doc:
            # Check for multiple blocks side-by-side indicating columns
            blocks = page.get_text("blocks")
            # A simple heuristic: if multiple blocks have significantly different x0 coordinates, it might be multi-column
            x0_values = [b[0] for b in blocks if b[4].strip()]
            if len(set(round(x/50) for x in x0_values)) > 3:
                parseability_score -= 8
                deductions.append("-8%: Multi-column layout detected. This may not parse correctly in legacy ATS like Workday or iCIMS.")
                break
    except Exception as e:
        logger.error("PDF Parsing failed: %s", e)
        parseability_score -= 20
        deductions.append("-20%: Failed to extract text from PDF (Unreadable layout or flattened image).")

    # Deductions from skipped fields
    for field in skipped:
        parseability_score -= 5
        deductions.append(f"-5%: Missing recommended field: {field.get('name', 'Unknown')}. Typically expected for this role.")

    # Job Match Score
    job_match_score = 100
    if job_description.strip():
        # Extracted text
        text = ""
        try:
            for page in pdf_doc:
                text += page.get_text()
        except:
            pass
        
        # We can call Gemini to cross-check overlap and suggest improvements
        sys_prompt = "You are an expert ATS Job Match Scorer. Compare the resume text to the job description."
        user_prompt = (
            f"RESUME:\n{text}\n\nJOB DESCRIPTION:\n{job_description}\n\n"
            "Analyze the resume strictly. Identify missing mandatory keywords from the JD.\n"
            "Also, provide 2-3 concrete improvement suggestions (e.g., 'Quantify your experience in X', 'Add a section for Y').\n"
            "Respond with ONLY JSON: { \"score\": 85, \"missing_keywords\": [\"Kubernetes\"], \"improvements\": [\"Quantify cloud migration impact\"] }"
        )
        try:
            match_res = _call_gemini(sys_prompt, user_prompt)
            job_match_score = match_res.get("score", 100)
            for kw in match_res.get("missing_keywords", []):
                job_match_score -= 2
                deductions.append(f"-2%: Missing keyword '{kw}' found in the job description.")
            for imp in match_res.get("improvements", []):
                improvements.append(imp)
        except:
            pass
    else:
        job_match_score = None # Not applicable

    parseability_score = max(0, parseability_score)
    if job_match_score is not None:
        job_match_score = max(0, job_match_score)

    return {
        "parseability_score": parseability_score,
        "job_match_score": job_match_score,
        "deductions": deductions,
        "improvements": improvements
    }

class ScoreBuilderRequest(BaseModel):
    resume_text: str
    job_description: str

@router.post("/score_builder")
async def score_builder_resume(
    req: ScoreBuilderRequest,
    user_and_client=Depends(get_user_and_scoped_client)
):
    # Templates generated by our builder are guaranteed ATS-parseable
    parseability_score = 100
    deductions = []
    improvements = []
    job_match_score = 100

    if req.job_description.strip():
        sys_prompt = "You are an expert ATS Job Match Scorer. Compare the resume text to the job description."
        user_prompt = (
            f"RESUME:\n{req.resume_text}\n\nJOB DESCRIPTION:\n{req.job_description}\n\n"
            "Analyze the resume strictly. Identify missing mandatory keywords from the JD.\n"
            "Also, provide 2-3 concrete improvement suggestions.\n"
            "Respond with ONLY JSON: { \"score\": 85, \"missing_keywords\": [\"Kubernetes\"], \"improvements\": [\"Quantify cloud migration impact\"] }"
        )
        try:
            match_res = _call_gemini(sys_prompt, user_prompt)
            job_match_score = match_res.get("score", 100)
            for kw in match_res.get("missing_keywords", []):
                job_match_score -= 2
                deductions.append(f"-2%: Missing keyword '{kw}' found in the job description.")
            for imp in match_res.get("improvements", []):
                improvements.append(imp)
        except:
            pass
    else:
        job_match_score = None

    if job_match_score is not None:
        job_match_score = max(0, job_match_score)

    return {
        "parseability_score": parseability_score,
        "job_match_score": job_match_score,
        "deductions": deductions,
        "improvements": improvements
    }

@router.post("/parse")
async def parse_resume(
    file: UploadFile = File(...),
    target_role: Optional[str] = Form(""),
    job_description: Optional[str] = Form(""),
    user_and_client=Depends(get_user_and_scoped_client)
):
    current_user, _ = user_and_client
    langfuse_context.update_current_observation(user_id=current_user.id)
    
    file_bytes = await file.read()
    
    # Extract text
    text = ""
    try:
        if file.filename.lower().endswith(".pdf"):
            pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in pdf_doc:
                text += page.get_text() + "\n"
        elif file.filename.lower().endswith(".docx"):
            doc = Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            text = file_bytes.decode("utf-8", errors="ignore")
    except Exception as e:
        logger.error("Failed to parse file text: %s", e)
        raise HTTPException(status_code=400, detail="Could not extract text from file.")

    sys_prompt = (
        "You are an expert resume parser and writer. Extract the user's information from the text "
        "into a structured JSON format that exactly matches the following schema:\n"
        "{\n"
        '  "contact": { "fullName": "", "email": "", "phone": "", "location": "", "linkedin": "", "github": "" },\n'
        '  "summary": "",\n'
        '  "experience": [ { "title": "", "company": "", "dates": "", "location": "", "bullets": ["", ""] } ],\n'
        '  "education": [ { "degree": "", "school": "", "dates": "", "gpa": "", "location": "" } ],\n'
        '  "skills": [ { "category": "Languages", "items": ["", ""] } ],\n'
        '  "projects": [ { "title": "", "link": "", "dates": "", "bullets": ["", ""] } ],\n'
        '  "certifications": [ { "name": "", "issuer": "", "date": "", "summary": "" } ]\n'
        "}\n"
        "BE EXTREMELY THOROUGH. Extract every single bullet point, project, and certification you can find in the text. Do not summarize or omit details from the experience bullets."
        "\nReturn ONLY the valid JSON object without any markdown wrapping or extra text."
    )
    if target_role or job_description:
        sys_prompt += f" Additionally, intelligently tailor and optimize the extracted resume content specifically for the target role: {target_role} and job description: {job_description}. Ensure professional tone and ATS compatibility while keeping all extracted details."

    user_prompt = f"RESUME TEXT:\n{text[:20000]}"

    try:
        result = _run_ai_generation(sys_prompt, user_prompt)
        langfuse_context.flush()
        return result
    except Exception as e:
        logger.error(f"AI parsing failed: {e}")
        raise HTTPException(status_code=502, detail="Failed to parse resume with AI.")

@router.post("/export_docx")
async def export_docx(state: ResumeState, user_and_client=Depends(get_user_and_scoped_client)):
    doc = Document()
    
    # Simple DOCX mapping
    data = state.resume_data
    doc.add_heading(data.get("contact", {}).get("fullName", "Resume"), 0)
    doc.add_paragraph(f"{data.get('contact', {}).get('email', '')} | {data.get('contact', {}).get('phone', '')}")
    
    if data.get("summary"):
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(data["summary"])
        
    if data.get("experience"):
        doc.add_heading("Experience", level=1)
        for exp in data["experience"]:
            doc.add_heading(f"{exp.get('title', '')} at {exp.get('company', '')}", level=2)
            doc.add_paragraph(exp.get("dates", ""))
            for b in exp.get("bullets", []):
                doc.add_paragraph(b, style='List Bullet')
                
    if data.get("education"):
        doc.add_heading("Education", level=1)
        for ed in data["education"]:
            doc.add_paragraph(ed)
            
    if data.get("skills"):
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(data["skills"]))

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    
    safe_name = (data.get("contact", {}).get("fullName", "resume")).replace(" ", "_")
    filename = f"{safe_name}_Builder.docx"
    
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
